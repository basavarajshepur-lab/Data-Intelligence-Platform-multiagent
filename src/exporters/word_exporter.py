"""Word exporter — editable data catalogue document using python-docx.

Produces a styled .docx ready for data steward review and sign-off.
Sensitivity levels are colour-coded in the field table.
"""

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..schema import DatasetMetadata, SensitivityLevel

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF8, 0xFA, 0xFC)
MID_GRAY = RGBColor(0x64, 0x74, 0x8B)
BLACK = RGBColor(0x0F, 0x17, 0x2A)

SENSITIVITY_RGB = {
    "public":       RGBColor(0x2D, 0x9E, 0x47),
    "internal":     RGBColor(0x2E, 0x86, 0xAB),
    "confidential": RGBColor(0xF1, 0x8F, 0x01),
    "restricted":   RGBColor(0xD1, 0x32, 0x32),
    "secret":       RGBColor(0x8B, 0x00, 0x00),
}

SENSITIVITY_HEX = {
    "public":       "2D9E47",
    "internal":     "2E86AB",
    "confidential": "F18F01",
    "restricted":   "D13232",
    "secret":       "8B0000",
}


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 8,
                   color: RGBColor = BLACK, align: str = "left"):
    para = cell.paragraphs[0]
    para.clear()
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = alignment_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _add_header_row(table, headers: list[str], col_widths: list[float]):
    row = table.rows[0]
    for i, (header, width) in enumerate(zip(headers, col_widths)):
        cell = row.cells[i]
        cell.width = Cm(width)
        _set_cell_bg(cell, "1E3A5F")
        _set_cell_text(cell, header, bold=True, size=8, color=WHITE, align="center")


def _heading(doc: Document, text: str, level: int = 1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(12 if level == 1 else 10)
    return para


def _kv_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        cells = table.rows[i].cells
        _set_cell_bg(cells[0], "F0F4F8")
        _set_cell_text(cells[0], key, bold=True, size=8)
        _set_cell_text(cells[1], val or "—", size=8)
        cells[0].width = Cm(5)
        cells[1].width = Cm(12)
    doc.add_paragraph()


def export(metadata: DatasetMetadata) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ---- Cover --------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(metadata.dataset_name)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    run2 = sub.add_run(
        f"Metadata Catalogue Entry  |  {metadata.data_domain.value.title()} Domain  |  "
        f"v{metadata.version}  |  Generated {datetime.now().strftime('%d %b %Y')}"
    )
    run2.font.size = Pt(9)
    run2.font.color.rgb = MID_GRAY

    # Classification banner (coloured paragraph)
    cls_para = doc.add_paragraph()
    cls_run = cls_para.add_run(
        f"  DATA CLASSIFICATION: {metadata.classification.value.upper()}  "
    )
    cls_run.bold = True
    cls_run.font.size = Pt(9)
    cls_run.font.color.rgb = WHITE
    shade_hex = SENSITIVITY_HEX.get(metadata.classification.value, "2E86AB")
    _set_cell_bg(cls_para._p, shade_hex)  # Can't shade para directly, use table below instead

    doc.add_paragraph()

    # ---- Dataset Overview --------------------------------------------
    _heading(doc, "1. Dataset Overview")
    _kv_table(doc, [
        ("Dataset Name", metadata.dataset_name),
        ("Version", metadata.version),
        ("Data Domain", metadata.data_domain.value.title()),
        ("Sub-Domain", metadata.sub_domain or "—"),
        ("Classification", metadata.classification.value.upper()),
        ("Classification Rationale", metadata.data_classification_rationale),
        ("Source System", metadata.source_system or "Not specified"),
        ("Data Steward", metadata.data_steward or "Not specified"),
        ("Data Owner", metadata.data_owner or "Not specified"),
        ("Schema Version", metadata.schema_version),
        ("Generated By", metadata.generated_by),
        ("Generated At", metadata.generated_at),
    ])

    _heading(doc, "Description", level=2)
    doc.add_paragraph(metadata.description).runs[0].font.size = Pt(9)

    _heading(doc, "Business Context", level=2)
    doc.add_paragraph(metadata.business_context).runs[0].font.size = Pt(9)

    _heading(doc, "Usage Guidance", level=2)
    doc.add_paragraph(metadata.usage_guidance).runs[0].font.size = Pt(9)

    if metadata.known_limitations:
        _heading(doc, "Known Limitations", level=2)
        doc.add_paragraph(metadata.known_limitations).runs[0].font.size = Pt(9)

    doc.add_page_break()

    # ---- Field Inventory -------------------------------------------
    _heading(doc, f"2. Field Inventory  ({len(metadata.fields)} fields)")

    col_widths = [3.5, 2.0, 1.5, 2.5, 7.5]
    headers = ["Field Name", "Data Type", "PII", "Sensitivity", "Description"]
    table = doc.add_table(rows=1 + len(metadata.fields), cols=5)
    table.style = "Table Grid"
    _add_header_row(table, headers, col_widths)

    for i, field in enumerate(metadata.fields, start=1):
        row = table.rows[i]
        bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"

        # Field name
        row.cells[0].width = Cm(col_widths[0])
        _set_cell_bg(row.cells[0], bg)
        _set_cell_text(row.cells[0], field.name, bold=True, size=8)

        # Data type
        row.cells[1].width = Cm(col_widths[1])
        _set_cell_bg(row.cells[1], bg)
        _set_cell_text(row.cells[1], field.data_type.value, size=8, align="center")

        # PII
        row.cells[2].width = Cm(col_widths[2])
        if field.is_pii:
            _set_cell_bg(row.cells[2], "FFEBEB")
            _set_cell_text(row.cells[2], field.pii_type.value if field.pii_type else "Yes",
                           bold=True, size=7, color=RGBColor(0x8B, 0, 0), align="center")
        else:
            _set_cell_bg(row.cells[2], bg)
            _set_cell_text(row.cells[2], "", size=8)

        # Sensitivity
        row.cells[3].width = Cm(col_widths[3])
        sens_hex = SENSITIVITY_HEX.get(field.sensitivity_level.value, "2E86AB")
        _set_cell_bg(row.cells[3], sens_hex)
        _set_cell_text(row.cells[3], field.sensitivity_level.value.upper(),
                       bold=True, size=7, color=WHITE, align="center")

        # Description
        row.cells[4].width = Cm(col_widths[4])
        _set_cell_bg(row.cells[4], bg)
        _set_cell_text(row.cells[4], field.description[:200], size=8)

    doc.add_paragraph()
    doc.add_page_break()

    # ---- Detailed Field Metadata -----------------------------------
    _heading(doc, "3. Detailed Field Metadata")

    for field in metadata.fields:
        _heading(doc, field.name, level=2)
        detail_rows = [
            ("Display Name", field.display_name),
            ("Data Type", field.data_type.value),
            ("Format", field.format or "—"),
            ("Is PII", "Yes" if field.is_pii else "No"),
            ("PII Type", field.pii_type.value if field.pii_type else "—"),
            ("Sensitivity", field.sensitivity_level.value.upper()),
            ("Is Key Field", "Yes" if field.is_key_field else "No"),
            ("Nullable", "Yes" if field.constraints.nullable else "No"),
            ("Unique", "Yes" if field.constraints.unique else "No"),
            ("Pattern", field.constraints.pattern or "—"),
            ("Foreign Key", field.constraints.foreign_key_ref or "—"),
            ("Description", field.description),
            ("Business Context", field.business_context),
            ("Usage Guidance", field.usage_guidance),
            ("Business Rules", field.business_rules or "—"),
            ("Data Lineage", field.data_lineage or "—"),
            ("Quality Notes", field.quality_notes or "—"),
            ("Tags", ", ".join(field.tags) or "—"),
        ]
        if field.guardrail_applied:
            detail_rows.append(("Guardrail Applied", field.guardrail_applied))
        _kv_table(doc, detail_rows)

    doc.add_page_break()

    # ---- PII Register ----------------------------------------------
    pii_fields = [f for f in metadata.fields if f.is_pii]
    if pii_fields:
        _heading(doc, f"4. PII Field Register  ({len(pii_fields)} fields)")
        pii_table = doc.add_table(rows=1 + len(pii_fields), cols=4)
        pii_table.style = "Table Grid"
        _add_header_row(pii_table, ["Field Name", "PII Type", "Sensitivity", "Usage Guidance"], [4, 3.5, 3, 6.5])

        for i, field in enumerate(sorted(pii_fields, key=lambda f: f.sensitivity_level.rank, reverse=True), start=1):
            bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"
            row = pii_table.rows[i]
            _set_cell_bg(row.cells[0], bg)
            _set_cell_text(row.cells[0], field.name, bold=True, size=8)
            _set_cell_bg(row.cells[1], bg)
            _set_cell_text(row.cells[1], field.pii_type.value if field.pii_type else "—", size=8)
            sens_hex = SENSITIVITY_HEX.get(field.sensitivity_level.value, "2E86AB")
            _set_cell_bg(row.cells[2], sens_hex)
            _set_cell_text(row.cells[2], field.sensitivity_level.value.upper(), bold=True, size=7, color=WHITE, align="center")
            _set_cell_bg(row.cells[3], bg)
            _set_cell_text(row.cells[3], field.usage_guidance[:120], size=7)

        doc.add_paragraph()
        doc.add_page_break()

    # ---- Compliance -----------------------------------------------
    _heading(doc, f"{'5' if pii_fields else '4'}. Compliance & Regulatory")
    c = metadata.compliance
    _kv_table(doc, [
        ("GDPR Applicable", "Yes" if c.gdpr_applicable else "No"),
        ("UK GDPR Applicable", "Yes" if c.uk_gdpr_applicable else "No"),
        ("Regulatory Frameworks", ", ".join(f.value for f in c.regulatory_frameworks) or "None"),
        ("Retention Period", c.retention_period or "Not specified"),
        ("Data Residency", c.data_residency_requirements or "Not specified"),
        ("Cross-Border Transfer Restrictions", "Yes" if c.cross_border_transfer_restrictions else "No"),
        ("Consent Required", "Yes" if c.consent_required else "No"),
        ("Right to Erasure Applicable", "Yes" if c.right_to_erasure_applicable else "No"),
        ("Lawful Basis", c.lawful_basis or "Not specified"),
    ])

    # ---- Quality Report -------------------------------------------
    qs = metadata.quality_score
    if qs:
        doc.add_page_break()
        n = 6 if pii_fields else 5
        _heading(doc, f"{n}. Quality Evaluation Report")

        overall_para = doc.add_paragraph()
        run = overall_para.add_run(
            f"Overall Score: {qs.overall_score:.1f}/100   {'PASSED' if qs.passed else 'FAILED'}"
        )
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2D, 0x9E, 0x47) if qs.passed else RGBColor(0xD1, 0x32, 0x32)
        doc.add_paragraph()

        q_table = doc.add_table(rows=1 + 5, cols=4)
        q_table.style = "Table Grid"
        _add_header_row(q_table, ["Dimension", "Score", "Weight", "Status"], [7, 2.5, 2, 2.5])

        dims = [
            ("Completeness", qs.completeness, "30%"),
            ("PII Detection", qs.pii_detection, "25%"),
            ("Type Consistency", qs.type_consistency, "20%"),
            ("Banking Standards (BCBS 239)", qs.banking_standards, "15%"),
            ("Sensitivity Consistency", qs.sensitivity_consistency, "10%"),
        ]
        for i, (name, dim, weight) in enumerate(dims, start=1):
            bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"
            row = q_table.rows[i]
            _set_cell_bg(row.cells[0], bg)
            _set_cell_text(row.cells[0], name, size=8)
            _set_cell_bg(row.cells[1], bg)
            _set_cell_text(row.cells[1], f"{dim.score:.1f}/100", size=8, align="center")
            _set_cell_bg(row.cells[2], bg)
            _set_cell_text(row.cells[2], weight, size=8, align="center")
            status_hex = "2D9E47" if dim.passed else "D13232"
            _set_cell_bg(row.cells[3], status_hex)
            _set_cell_text(row.cells[3], "PASS" if dim.passed else "FAIL",
                           bold=True, size=8, color=WHITE, align="center")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
